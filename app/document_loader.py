import csv
import re
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

from langchain_core.documents import Document
from pypdf import PdfReader


FALLBACK_CHUNK_SIZE = 1000
FALLBACK_CHUNK_OVERLAP = 150
ALLOWED_FILE_EXTENSIONS = {".pdf", ".txt", ".md", ".docx", ".csv"}


@dataclass
class ProcessingResult:
    documents: list[Document]
    parser_used: str
    warnings: list[str] = field(default_factory=list)
    detected_extension: str | None = None


def load_text_as_documents(
    text: str,
    source_id: str,
    filename: str | None,
    doc_type: str,
) -> list[Document]:
    documents = _documents_from_text(
        text=text,
        source_id=source_id,
        filename=filename,
        doc_type=doc_type,
        page_number=None,
        parser_used="plain_text",
    )
    return documents


def load_file_as_documents(
    file_path: str,
    source_id: str,
    filename: str,
    doc_type: str,
) -> ProcessingResult:
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".txt":
        text = _read_text_file(path)
        documents = _documents_from_text(
            text,
            source_id,
            filename,
            doc_type,
            page_number=None,
            parser_used="plain_text",
        )
        return ProcessingResult(documents, "plain_text", detected_extension=suffix)

    if suffix == ".md":
        text = _read_text_file(path)
        documents = _documents_from_markdown(text, source_id, filename, doc_type)
        return ProcessingResult(documents, "markdown_text", detected_extension=suffix)

    if suffix == ".csv":
        documents, warnings = _load_csv(path, source_id, filename, doc_type)
        return ProcessingResult(documents, "csv", warnings, detected_extension=suffix)

    if suffix == ".docx":
        documents = _load_docx(path, source_id, filename, doc_type)
        return ProcessingResult(documents, "docx", detected_extension=suffix)

    docling_error: Exception | None = None
    try:
        documents = _load_with_docling(path, source_id, filename, doc_type)
        if documents:
            warnings = _pdf_warnings(documents) if suffix == ".pdf" else []
            return ProcessingResult(documents, "docling", warnings, detected_extension=suffix)
    except Exception as exc:
        docling_error = exc

    if suffix == ".pdf":
        try:
            documents = _load_pdf_fallback(path, source_id, filename, doc_type)
            return ProcessingResult(documents, "pypdf", detected_extension=suffix)
        except Exception as exc:
            raise ValueError(
                f"Could not parse PDF with Docling or pypdf fallback: {exc}"
            ) from exc

    supported = ", ".join(sorted(ALLOWED_FILE_EXTENSIONS))
    if docling_error is not None:
        raise ValueError(
            f"Could not parse {filename} with Docling. Supported fallback formats are {supported}. "
            f"Docling error: {docling_error}"
        ) from docling_error

    raise ValueError(f"Unsupported file type '{suffix or 'unknown'}'. Supported formats: {supported}.")


def _load_with_docling(
    path: Path,
    source_id: str,
    filename: str,
    doc_type: str,
) -> list[Document]:
    from docling.document_converter import DocumentConverter

    converter = DocumentConverter()
    result = converter.convert(str(path))
    docling_document = getattr(result, "document", None)
    if docling_document is None:
        raise ValueError("Docling did not return a document object.")

    chunks = _chunk_docling_document(docling_document)
    if chunks:
        documents: list[Document] = []
        for index, chunk in enumerate(chunks):
            text = _extract_chunk_text(chunk)
            if not text:
                continue
            metadata = _base_metadata(source_id, filename, doc_type, index)
            metadata["parser_used"] = "docling"
            metadata["chunk_char_count"] = len(text)
            page_number = _extract_page_number(chunk)
            if page_number is not None:
                metadata["page_number"] = page_number
            documents.append(Document(page_content=text, metadata=metadata))
        if documents:
            return documents

    exported_text = _export_docling_text(docling_document)
    return _documents_from_text(
        exported_text,
        source_id,
        filename,
        doc_type,
        page_number=None,
        parser_used="docling",
    )


def _chunk_docling_document(docling_document: Any) -> list[Any]:
    try:
        from docling.chunking import HybridChunker
    except Exception:
        try:
            from docling.chunking.hybrid_chunker import HybridChunker
        except Exception:
            return []

    chunker = HybridChunker()
    try:
        return list(chunker.chunk(dl_doc=docling_document))
    except TypeError:
        return list(chunker.chunk(docling_document))


def _extract_chunk_text(chunk: Any) -> str:
    if isinstance(chunk, str):
        return chunk.strip()
    for attr in ("text", "page_content"):
        value = getattr(chunk, attr, None)
        if isinstance(value, str) and value.strip():
            return value.strip()
    export_text = getattr(chunk, "export_text", None)
    if callable(export_text):
        value = export_text()
        if isinstance(value, str):
            return value.strip()
    return str(chunk).strip()


def _extract_page_number(chunk: Any) -> int | None:
    metadata = getattr(chunk, "metadata", None)
    if isinstance(metadata, dict):
        for key in ("page_number", "page", "page_no"):
            value = metadata.get(key)
            if isinstance(value, int):
                return value
    return None


def _export_docling_text(docling_document: Any) -> str:
    for method_name in ("export_to_markdown", "export_to_text"):
        method = getattr(docling_document, method_name, None)
        if callable(method):
            text = method()
            if isinstance(text, str) and text.strip():
                return text
    raise ValueError("Docling parsed the document but could not export text.")


def _load_pdf_fallback(
    path: Path,
    source_id: str,
    filename: str,
    doc_type: str,
) -> list[Document]:
    reader = PdfReader(str(path))
    documents: list[Document] = []
    chunk_index = 0
    for page_index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        for chunk in _split_text(text):
            metadata = _base_metadata(source_id, filename, doc_type, chunk_index)
            metadata["page_number"] = page_index
            metadata["parser_used"] = "pypdf"
            metadata["chunk_char_count"] = len(chunk)
            documents.append(Document(page_content=chunk, metadata=metadata))
            chunk_index += 1
    if not documents:
        raise ValueError("No text could be extracted from the PDF.")
    return documents


def _documents_from_text(
    text: str,
    source_id: str,
    filename: str | None,
    doc_type: str,
    page_number: int | None,
    parser_used: str,
    section_title: str | None = None,
) -> list[Document]:
    documents: list[Document] = []
    for index, chunk in enumerate(_split_text(text)):
        metadata = _base_metadata(source_id, filename, doc_type, index)
        if page_number is not None:
            metadata["page_number"] = page_number
        metadata["parser_used"] = parser_used
        metadata["chunk_char_count"] = len(chunk)
        if section_title:
            metadata["section_title"] = section_title
        documents.append(Document(page_content=chunk, metadata=metadata))
    return documents


def _documents_from_markdown(
    text: str,
    source_id: str,
    filename: str,
    doc_type: str,
) -> list[Document]:
    documents: list[Document] = []
    current_section = None
    chunk_index = 0
    for section_title, section_text in _markdown_sections(text):
        section_documents = _documents_from_text(
            section_text,
            source_id,
            filename,
            doc_type,
            page_number=None,
            parser_used="markdown_text",
            section_title=section_title,
        )
        for document in section_documents:
            document.metadata["chunk_index"] = chunk_index
            chunk_index += 1
            documents.append(document)
        current_section = section_title

    if documents:
        return documents
    return _documents_from_text(
        text,
        source_id,
        filename,
        doc_type,
        page_number=None,
        parser_used="markdown_text",
        section_title=current_section,
    )


def _markdown_sections(text: str) -> list[tuple[str | None, str]]:
    sections: list[tuple[str | None, list[str]]] = []
    current_title: str | None = None
    current_lines: list[str] = []
    heading_pattern = re.compile(r"^\s{0,3}(#{1,6})\s+(.+?)\s*$")

    for line in text.splitlines():
        match = heading_pattern.match(line)
        if match and current_lines:
            sections.append((current_title, current_lines))
            current_lines = [line]
            current_title = match.group(2).strip()
        else:
            if match:
                current_title = match.group(2).strip()
            current_lines.append(line)

    if current_lines:
        sections.append((current_title, current_lines))
    return [(title, "\n".join(lines).strip()) for title, lines in sections if "\n".join(lines).strip()]


def _load_csv(
    path: Path,
    source_id: str,
    filename: str,
    doc_type: str,
) -> tuple[list[Document], list[str]]:
    warnings: list[str] = []
    documents: list[Document] = []
    chunk_index = 0
    skipped_empty_rows = 0

    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.reader(handle)
        rows = list(reader)

    if not rows:
        return [], []

    header = [cell.strip() for cell in rows[0]]
    has_header = any(header)
    for row_number, row in enumerate(rows[1:] if has_header else rows, start=2 if has_header else 1):
        cells = [cell.strip() for cell in row]
        if not any(cells):
            skipped_empty_rows += 1
            continue

        text = _format_csv_row(header if has_header else [], cells, row_number)
        for chunk in _split_text(text):
            metadata = _base_metadata(source_id, filename, doc_type, chunk_index)
            metadata["parser_used"] = "csv"
            metadata["row_number"] = row_number
            metadata["section_title"] = "CSV rows"
            metadata["chunk_char_count"] = len(chunk)
            documents.append(Document(page_content=chunk, metadata=metadata))
            chunk_index += 1

    if skipped_empty_rows:
        warnings.append("CSV contained empty rows that were skipped.")
    return documents, warnings


def _format_csv_row(header: list[str], cells: list[str], row_number: int) -> str:
    if header:
        pairs = []
        for index, cell in enumerate(cells):
            column = header[index] if index < len(header) and header[index] else f"column_{index + 1}"
            pairs.append(f"{column}: {cell}")
        return f"CSV row {row_number}\n" + "\n".join(pairs)
    return f"CSV row {row_number}\n" + ", ".join(cells)


def _load_docx(
    path: Path,
    source_id: str,
    filename: str,
    doc_type: str,
) -> list[Document]:
    try:
        text = _read_docx_with_python_docx(path)
    except Exception:
        text = _read_docx_with_stdlib(path)
    if not text.strip():
        raise ValueError("No text could be extracted from the DOCX file.")
    return _documents_from_text(
        text,
        source_id,
        filename,
        doc_type,
        page_number=None,
        parser_used="docx",
    )


def _read_docx_with_python_docx(path: Path) -> str:
    from docx import Document as DocxDocument

    docx_document = DocxDocument(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in docx_document.paragraphs if paragraph.text.strip()]
    return "\n\n".join(paragraphs)


def _read_docx_with_stdlib(path: Path) -> str:
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(path) as archive:
        xml_bytes = archive.read("word/document.xml")
    root = ElementTree.fromstring(xml_bytes)
    paragraphs = []
    for paragraph in root.findall(".//w:p", namespace):
        texts = [node.text or "" for node in paragraph.findall(".//w:t", namespace)]
        paragraph_text = "".join(texts).strip()
        if paragraph_text:
            paragraphs.append(paragraph_text)
    return "\n\n".join(paragraphs)


def _pdf_warnings(documents: list[Document]) -> list[str]:
    if not any(document.metadata.get("page_number") is not None for document in documents):
        return ["No page metadata found."]
    return []


def _split_text(text: str) -> list[str]:
    normalized = text.strip()
    if not normalized:
        return []

    chunks: list[str] = []
    start = 0
    while start < len(normalized):
        end = min(start + FALLBACK_CHUNK_SIZE, len(normalized))
        chunk = normalized[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(normalized):
            break
        start = max(end - FALLBACK_CHUNK_OVERLAP, start + 1)
    return chunks


def _base_metadata(
    source_id: str,
    filename: str | None,
    doc_type: str,
    chunk_index: int,
) -> dict[str, Any]:
    return {
        "source_id": source_id,
        "filename": filename,
        "doc_type": doc_type,
        "chunk_index": chunk_index,
    }


def _read_text_file(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="latin-1")
